"""Text extractors for different file types."""

import hashlib
from abc import ABC, abstractmethod
from datetime import datetime
from pathlib import Path

from ..utils.logging import get_logger

logger = get_logger(__name__)


class ExtractionError(Exception):
    """Raised when text extraction fails."""

    pass


class BaseExtractor(ABC):
    """Base class for text extractors."""

    @property
    @abstractmethod
    def supported_extensions(self) -> set[str]:
        """Return set of supported file extensions (lowercase, with dot)."""
        pass

    @abstractmethod
    def extract(self, file_path: Path) -> str:
        """
        Extract text content from a file.

        Args:
            file_path: Path to the file

        Returns:
            Extracted text content

        Raises:
            ExtractionError: If extraction fails
        """
        pass

    def can_handle(self, file_path: Path) -> bool:
        """Check if this extractor can handle the given file."""
        return file_path.suffix.lower() in self.supported_extensions


class PlainTextExtractor(BaseExtractor):
    """Extractor for plain text files (.txt, .md)."""

    @property
    def supported_extensions(self) -> set[str]:
        return {".txt", ".md"}

    def extract(self, file_path: Path) -> str:
        """Extract text from plain text files."""
        try:
            # Try UTF-8 first, then fall back to other encodings
            encodings = ["utf-8", "utf-8-sig", "latin-1", "cp1252"]

            for encoding in encodings:
                try:
                    with open(file_path, "r", encoding=encoding) as f:
                        content = f.read()
                    logger.debug(f"Extracted text from {file_path} using {encoding}")
                    return content
                except UnicodeDecodeError:
                    continue

            raise ExtractionError(f"Could not decode file with any supported encoding: {file_path}")

        except OSError as e:
            raise ExtractionError(f"Could not read file {file_path}: {e}") from e


class PDFExtractor(BaseExtractor):
    """Extractor for PDF files using PyMuPDF (fitz)."""

    @property
    def supported_extensions(self) -> set[str]:
        return {".pdf"}

    def extract(self, file_path: Path) -> str:
        """Extract text from PDF files."""
        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise ExtractionError(
                "PyMuPDF (fitz) is not installed. Install with: pip install pymupdf"
            )

        try:
            doc = fitz.open(file_path)
            text_parts: list[str] = []

            for page_num, page in enumerate(doc):
                page_text = page.get_text()
                if page_text.strip():
                    text_parts.append(page_text)
                    logger.debug(f"Extracted {len(page_text)} chars from page {page_num + 1}")

            doc.close()

            full_text = "\n\n".join(text_parts)
            logger.debug(f"Extracted {len(full_text)} total chars from PDF: {file_path}")

            return full_text

        except Exception as e:
            raise ExtractionError(f"Failed to extract text from PDF {file_path}: {e}") from e


class DOCXExtractor(BaseExtractor):
    """Extractor for DOCX files using python-docx."""

    @property
    def supported_extensions(self) -> set[str]:
        return {".docx"}

    def extract(self, file_path: Path) -> str:
        """Extract text from DOCX files."""
        try:
            from docx import Document
        except ImportError:
            raise ExtractionError(
                "python-docx is not installed. Install with: pip install python-docx"
            )

        try:
            doc = Document(file_path)
            text_parts: list[str] = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)

            full_text = "\n".join(text_parts)
            logger.debug(f"Extracted {len(full_text)} chars from DOCX: {file_path}")

            return full_text

        except Exception as e:
            raise ExtractionError(f"Failed to extract text from DOCX {file_path}: {e}") from e


# Registry of all available extractors
EXTRACTORS: list[BaseExtractor] = [
    PlainTextExtractor(),
    PDFExtractor(),
    DOCXExtractor(),
]


def get_extractor(file_path: Path) -> BaseExtractor | None:
    """
    Get the appropriate extractor for a file.

    Args:
        file_path: Path to the file

    Returns:
        Extractor instance if one is available, None otherwise
    """
    for extractor in EXTRACTORS:
        if extractor.can_handle(file_path):
            return extractor
    return None


def get_supported_extensions() -> set[str]:
    """Get all supported file extensions across all extractors."""
    extensions: set[str] = set()
    for extractor in EXTRACTORS:
        extensions.update(extractor.supported_extensions)
    return extensions
